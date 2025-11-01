from fastapi import APIRouter, Form, HTTPException
from fastapi.responses import FileResponse
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from bson import ObjectId
from app.database import cvs_collection, jobs_collection, tailored_cvs_collection
from app.services.ollama_client import generate_tailored_cv
from app.services.cv_profile_service import cv_profile_to_text
from app.routes.analysis_routes import analyze_cv_job  # reuse logic
import logging

router = APIRouter(prefix="/generate", tags=["Generate"])
UPLOAD_DIR = Path("generated_cvs")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

logger = logging.getLogger(__name__)

@router.post("/")
async def generate_cv(
    cv_id: str = Form(...), 
    job_id: str = Form(...),
) -> dict:
    """Generate a tailored CV if similarity is high enough"""
    try:
        cv = await cvs_collection.find_one({"_id": ObjectId(cv_id)})
        job = await jobs_collection.find_one({"_id": ObjectId(job_id)})

        if not cv or not job:
            raise HTTPException(status_code=404, detail="CV or Job not found")

        analysis = await analyze_cv_job(cv_id=cv_id, job_id=job_id)
        similarity = analysis["similarity"]

        if similarity < 0.60:
            return {"status": "skipped", "similarity": similarity, "message": "Low match score"}

        cv_text = cv.get("raw_text", "")
        job_text = job.get("job_text", "")
        
        if not cv_text or not job_text:
            raise HTTPException(status_code=400, detail="CV or Job text is empty")
        
        tailored_content = generate_tailored_cv(cv_text, job_text)
        generated_id = f"tailored_{cv_id[:8]}_{job_id[:8]}"
        output_file = UPLOAD_DIR / f"{generated_id}.docx"
        create_docx(tailored_content, output_file, job.get("title", "Job Position"))

        await tailored_cvs_collection.insert_one({
            "generated_id": generated_id,
            "cv_id": cv_id,
            "job_id": job_id,
            "similarity": similarity,
            "tailored_text": tailored_content,
            "created_at": datetime.utcnow()
        })

        return {
            "status": "generated",
            "similarity": similarity,
            "generated_id": generated_id,
            "snippet": tailored_content[:300] + "...",
            "download_url": f"/generate/download/{generated_id}"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating CV: {e}", exc_info=True)
        error_msg = str(e)
        if "Ollama" in error_msg or "ollama" in error_msg.lower():
            raise HTTPException(
                status_code=500, 
                detail=f"Erreur lors de la génération du CV avec Ollama: {error_msg}"
            )
        raise HTTPException(status_code=500, detail=f"Error generating CV: {error_msg}")

def create_docx(content: str, output_path: Path, job_title: str):
    """Create a professional, well-formatted DOCX CV file"""
    from docx.shared import Pt, RGBColor, Inches
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set margins (professional CV margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Parse content and create professional format
    lines = content.splitlines()
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        # Handle main headings (##)
        if line.startswith('##'):
            heading_text = line.lstrip('#').strip()
            # Skip empty headings
            if heading_text:
                # Add spacing before main sections (except first)
                if i > 0:
                    doc.add_paragraph()
                heading = doc.add_heading(heading_text, level=1)
                if heading.runs:
                    heading_format = heading.runs[0].font
                    heading_format.size = Pt(14)
                    heading_format.bold = True
                    heading_format.color = RGBColor(0, 51, 102)  # Dark blue
                    # Add underline to main section headings
                    heading.runs[0].underline = True
                
        # Handle subheadings (###)
        elif line.startswith('###'):
            heading_text = line.lstrip('#').strip()
            if heading_text:
                subheading = doc.add_heading(heading_text, level=2)
                if subheading.runs:
                    subheading_format = subheading.runs[0].font
                    subheading_format.size = Pt(12)
                    subheading_format.bold = True
                    subheading_format.color = RGBColor(0, 0, 0)  # Black
                
        # Handle bullet points
        elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
            bullet_text = line.lstrip('-•* ').strip()
            if bullet_text:
                p = doc.add_paragraph(bullet_text, style='List Bullet')
                if p.runs:
                    p_format = p.runs[0].font
                    p_format.size = Pt(11)
                # Adjust bullet indent
                p.paragraph_format.space_after = Pt(3)
                
        # Handle regular paragraphs with potential formatting
        else:
            # Check if line contains key information patterns
            p = doc.add_paragraph(line)
            if p.runs:
                p_format = p.runs[0].font
                if any(keyword in line.lower() for keyword in ['email:', 'phone:', 'linkedin:', 'github:', 'location:']):
                    p_format.size = Pt(10.5)
                elif '|' in line and any(keyword in line.lower() for keyword in ['-', 'to', 'present']):
                    # Likely a date range or location line
                    p_format.size = Pt(10)
                    p_format.italic = True
                    p_format.color = RGBColor(64, 64, 64)  # Gray
                else:
                    p_format.size = Pt(11)
            
            # Add small spacing after paragraphs
            p.paragraph_format.space_after = Pt(6)
        
        i += 1
    
    # Ensure document ends with proper spacing
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
    
    doc.save(output_path)

@router.get("/download/{generated_id}")
async def download_generated(generated_id: str):
    file_path = UPLOAD_DIR / f"{generated_id}.docx"
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Generated CV not found")
    return FileResponse(
        file_path,
        filename=f"CV_Personnalise_{generated_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
